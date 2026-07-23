# -*- coding: utf-8 -*-
"""生成 DocuMind RAG 项目最新技术复盘报告（Word 版）。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 颜色 ----------
ACCENT = RGBColor(0x1F, 0x4E, 0x79)      # 主色 深蓝
ACCENT_LIGHT = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HDR_FILL = "1F4E79"
ALT_FILL = "EAF1F8"
CODE_FILL = "F4F4F4"

doc = Document()

# ---------- 基础样式 ----------
normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def set_cjk(run, font="Microsoft YaHei"):
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_cjk(run)
        if level == 1:
            run.font.color.rgb = ACCENT
        elif level == 2:
            run.font.color.rgb = ACCENT_LIGHT
    return h

def para(text="", size=10.5, bold=False, color=None, align=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    set_cjk(run)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        set_cjk(r1)
        r2 = p.add_run(text)
        set_cjk(r2)
    else:
        r = p.add_run(text)
        set_cjk(r)
    return p

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    set_cjk(run)

def add_table(headers, rows, widths=None, header_fill=HDR_FILL, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=WHITE, size=font_size)
        shade_cell(hdr[i], header_fill)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), size=font_size)
            if r_idx % 2 == 1:
                shade_cell(cells[i], ALT_FILL)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 背景底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), CODE_FILL)
    pPr.append(shd)
    return p

# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()
para("DocuMind RAG 知识库问答系统", size=26, bold=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para("技术复盘与架构报告（最新版）", size=16, bold=True, color=ACCENT_LIGHT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
para("— 涵盖 LangChain 升级、5 项中期功能增强与全部安全修复 —",
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

info = doc.add_table(rows=5, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
info_rows = [
    ("项目名称", "DocuMind — 基于 RAG 的多文档知识库问答 Agent"),
    ("报告版本", "v2.0（整合全部升级后的最新状态）"),
    ("生成日期", "2026-07-23"),
    ("技术栈", "Streamlit · ChromaDB · LangGraph · PyMuPDF · OpenAI SDK"),
    ("代码规模", "16 个 src 模块 + 4 大类单元测试（36 用例全过）"),
]
for i, (k, v) in enumerate(info_rows):
    set_cell_text(info.rows[i].cells[0], k, bold=True, color=WHITE)
    shade_cell(info.rows[i].cells[0], HDR_FILL)
    set_cell_text(info.rows[i].cells[1], v)
    info.rows[i].cells[0].width = Inches(1.6)
    info.rows[i].cells[1].width = Inches(4.6)
doc.add_page_break()

# ============================================================
# 目录（Word 自动域，打开后右键更新即可）
# ============================================================
heading("目录", level=1)
para("（在 Word 中右键“更新域”可自动生成页码）", size=9, italic=True, color=GREY)
toc_p = doc.add_paragraph()
run = toc_p.add_run()
fldStart = OxmlElement("w:fldChar"); fldStart.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
fldText = OxmlElement("w:t"); fldText.text = "右键更新目录"
fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
run._r.append(fldStart); run._r.append(instr); run._r.append(fldSep)
run._r.append(fldText); run._r.append(fldEnd)
doc.add_page_break()

# ============================================================
# 第一章 项目概览
# ============================================================
heading("一、项目概览", level=1)
para("DocuMind 是一个本地部署的、面向多文档知识库的检索增强生成（RAG）问答系统，"
     "核心是一个能够“规划—调用工具—综合答案”的多步 Agent。用户上传 PDF/TXT/MD/DOCX/PPTX/HTML/CSV "
     "等文档构建知识库后，可通过对话方式提问，Agent 自动检索相关片段、必要时二次检索对比，"
     "并生成带引用来源的答案；涉及删除/重建等危险操作时会生成“待确认操作”，需用户在界面二次确认才执行。", space_after=8)

para("能力矩阵：", bold=True, space_after=2)
bullet("语义/关键词混合检索：向量召回 + BM25 词法召回，RRF 融合；", bold_prefix="检索：")
bullet("重排序：零依赖词法 Reranker，可选 BGE Cross-Encoder 高质量重排；", bold_prefix="精排：")
bullet("多步 Agent：LangGraph ReAct 框架，工具调用 + 全局轨迹 trace；", bold_prefix="推理：")
bullet("安全护栏：危险操作二次确认、步数上限、异常回退；", bold_prefix="安全：")
bullet("对话记忆：token 预算滑动窗口 + LLM 摘要压缩；", bold_prefix="记忆：")
bullet("流式输出、基于来源的引用、知识库可视化管理。", bold_prefix="体验：")

# ============================================================
# 第二章 系统架构
# ============================================================
heading("二、系统架构", level=1)
para("项目采用清晰的层次化架构，从上一版的 8 层演进为“UI → Agent → RAG → LLM/检索/重排 → "
     "文档处理 → 知识库 → 配置”的多层结构。各模块职责单一、依赖方向自上而下，几乎不反向耦合。",
     space_after=8)

arch = (
    "┌─────────────────────────────────────────────────────────────┐\n"
    "│  UI 层：app.py (Streamlit)  —— 对话、侧栏控制、轨迹可视化、流式    │\n"
    "└───────────────────────────────┬─────────────────────────────┘\n"
    "                                 ↓\n"
    "┌─────────────────────────────────────────────────────────────┐\n"
    "│  Agent 层：src/agent.py (LangGraph create_react_agent)         │\n"
    "│   工具：检索 / 列文档 / 查信息 / 请求删除 / 请求重建              │\n"
    "│   机制：RunnableConfig 注入 ToolContext · 步数上限 · 回退 · trace  │\n"
    "└───────┬───────────────┬───────────────┬───────────────────────┘\n"
    "        ↓               ↓               ↓\n"
    "┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐\n"
    "│ RAG 合成层    │ │ 检索层        │ │ 工具/记忆层                │\n"
    "│ rag_chain.py │ │ retriever.py │ │ tools.py / memory.py      │\n"
    "│ 检索-生成解耦 │ │ QueryRewrite  │ │ 危险操作二次确认 + 记忆管理  │\n"
    "│ 流式合成      │ │ Hybrid+BM25   │ │                           │\n"
    "└──────┬───────┘ └──────┬───────┘ └────────────┬─────────────┘\n"
    "       ↓                ↓                      ↓\n"
    "┌──────────────┐ ┌──────────────┐ ┌──────────────────────────┐\n"
    "│ 重排层        │ │ 向量存储      │ │ 文档处理层                 │\n"
    "│ reranker.py  │ │ vectorstore   │ │ document_loader / splitter │\n"
    "│ 词法/CrossEnc │ │ .py (Chroma)  │ │ PyMuPDF/DOCX/PPTX/HTML/CSV │\n"
    "└──────────────┘ └──────────────┘ └────────────┬─────────────┘\n"
    "                                   ↓            ↓\n"
    "┌──────────────────────────┐ ┌──────────────────────────────┐\n"
    "│ 知识库管理层：kb_operations │ │ 配置/LLM 层：config / llm       │\n"
    "│ 备份·重建·删除·校验         │ │ dotenv + OpenAI SDK + LangChain │\n"
    "└──────────────────────────┘ └──────────────────────────────┘\n"
)
code_block(arch)

para("数据流向（一次问答）：", bold=True, space_after=2)
bullet("用户提问 → app.py 调用 agent_answer()；")
bullet("Agent 经 LangGraph 规划，调用 retrieve_documents 工具；")
bullet("retriever 先做 Query Rewrite，再走 Hybrid（向量+BM25 RRF）召回候选池；")
bullet("候选经 reranker 重排到 top_k，返回 chunks 并写入 ToolContext；")
bullet("Agent 取得证据后由 rag_chain 专用合成器生成带引用的答案；")
bullet("危险操作仅生成 pending_operation，等待 UI 二次确认；全程写入 trace。")

# ============================================================
# 第三章 模块与技术栈映射
# ============================================================
heading("三、模块与技术栈映射", level=1)
para("下表列出当前 src/ 下全部 16 个模块及其职责与所用技术。", space_after=6)
add_table(
    ["模块", "职责", "关键技术"],
    [
        ["app.py", "Streamlit UI：对话、侧栏、轨迹、流式", "streamlit"],
        ["src/agent.py", "多步 Agent 编排（规划/工具/合成）", "langgraph / langchain-core"],
        ["src/rag_chain.py", "检索-生成解耦 + 流式合成", "openai / langchain"],
        ["src/retriever.py", "Query Rewrite + Hybrid 检索 + 重排接入", "chromadb / rank_bm25 / reranker"],
        ["src/reranker.py", "检索结果重排序", "自研词法 + 可选 CrossEncoder"],
        ["src/vectorstore.py", "Chroma 集合、查询、元数据过滤", "chromadb"],
        ["src/document_loader.py", "PDF/DOCX/PPTX/HTML/CSV/TXT 解析", "PyMuPDF / python-docx / pptx / bs4"],
        ["src/splitter.py", "段落感知切块", "自研"],
        ["src/tools.py", "5 个 Agent 工具实现", "vectorstore / kb_operations"],
        ["src/memory.py", "对话记忆（token 预算 + 摘要）", "自研 + llm"],
        ["src/llm.py", "OpenAI 客户端 + LangChain ChatModel 工厂", "openai / langchain-openai"],
        ["src/prompts.py", "全部 Prompt 模板", "自研"],
        ["src/kb_operations.py", "知识库构建/删除/重建/备份", "chromadb / path 安全"],
        ["src/utils.py", "公共去重等工具", "自研"],
        ["src/config.py", "配置、开关、启动校验", "python-dotenv"],
        ["tests/", "4 大类单元测试（36 用例）", "pytest"],
    ],
    widths=[1.5, 3.0, 2.2],
)

# ============================================================
# 第四章 核心技术选型与横向对比
# ============================================================
heading("四、核心技术选型与横向对比", level=1)
para("本章对每一项关键技术选型，对比 2–4 个同类方案，给出选择理由与适用边界。"
     "其中带 ★ 的为本轮升级新增或重构项。", space_after=6)

heading("4.1 前端 UI：Streamlit", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["Streamlit ★", "Python 原生、组件丰富、开发最快", "单线程、状态管理弱、难做复杂布局", "原型/内部工具最佳，沿用"],
        ["Gradio", "对话 UI 开箱即用、分享方便", "自定义 UI 受限", "偏 Demo 演示"],
        ["Chainlit", "专为 LLM/Agent 设计、内置会话/追踪", "生态较新、定制受限", "若重 Agent 观测可换"],
        ["原生 Web", "完全可控、可上生产", "开发成本高", "生产化时再考虑"],
    ],
    widths=[1.2, 2.2, 2.0, 1.6],
)
para("结论：当前阶段 Streamlit 性价比最高。若后续要做生产级高并发，需迁移到 FastAPI + 前端框架。",
     italic=True, color=GREY, space_after=8)

heading("4.2 向量数据库：ChromaDB", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["ChromaDB ★", "零配置、内置嵌入、metadata 过滤、本地持久化", "超百万级需迁移", "完全匹配当前规模"],
        ["FAISS", "极致性能、纯向量", "无 metadata、需自管索引", "仅纯向量场景"],
        ["Milvus", "分布式、超大规模", "部署重、运维成本高", "规模上来再换"],
        ["Qdrant", "性能强、过滤好、易部署", "需独立服务", "生产可选"],
    ],
    widths=[1.2, 2.2, 2.0, 1.6],
)
para("结论：Chroma 的“零配置 + 内置嵌入 + metadata 过滤”与本项目完美契合；当文档量进入百万级再评估 Milvus/Qdrant。",
     italic=True, color=GREY, space_after=8)

heading("4.3 文档解析：PyMuPDF（升级）", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["PyMuPDF(fitz) ★", "文本质量高、速度快、支持表格/版式", "依赖较重", "PDF 解析主力"],
        ["PyPDF", "极轻量", "提取质量一般、无表格", "回退方案"],
        ["pdfplumber", "表格提取强", "速度较慢", "表格场景可补"],
        ["Unstructured", "多格式一体、版面理解", "重依赖、黑盒", "企业级可评估"],
    ],
    widths=[1.4, 2.0, 2.0, 1.6],
)
para("结论：本项目已实现 PyMuPDF 优先、PyPDF 回退的解析链，并在 document_loader 中新增 DOCX/PPTX/HTML/CSV 支持，覆盖 7 种格式。",
     italic=True, color=GREY, space_after=8)

heading("4.4 Agent 框架：LangGraph（升级）", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["LangGraph ★", "官方 Agent 标准、状态图可控、易加护栏", "需理解图范式", "本轮升级采用"],
        ["LangChain AgentExecutor", "经典易用", "1.x 已移除", "不再适用"],
        ["自研循环(旧版)", "零依赖、完全透明", "需自维护、易出错", "已替换"],
        ["LlamaIndex Agent", "与索引深度整合", "框架绑定重", "非本项目选型"],
    ],
    widths=[1.4, 2.0, 2.0, 1.6],
)
para("结论：LangChain 1.x 移除旧 AgentExecutor，官方推荐 LangGraph 的 create_react_agent；"
     "本项目用它替代自研循环，并完整保留 5 工具、二次确认、回退、trace、检索-生成解耦。",
     italic=True, color=GREY, space_after=8)

heading("4.5 LLM 接口：OpenAI SDK + LangChain ChatModel（升级）", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["OpenAI SDK 薄封装 ★", "与模型解耦、可控 timeout/重试", "需自写流式", "聊天/合成主用"],
        ["LangChain ChatModel ★", "与 Agent 框架无缝、统一接口", "抽象层略厚", "Agent 调用用"],
        ["LangChain 全封装", "省事", "黑盒、难调优", "不采用"],
        ["LlamaIndex LLM", "生态全", "框架绑定", "不采用"],
    ],
    widths=[1.6, 1.9, 1.9, 1.6],
)
para("结论：本项目采用“双轨”——底层仍是 OpenAI SDK（含 timeout=60s、max_retries=2、流式），"
     "Agent 层用 LangChain ChatOpenAI 工厂，二者共享同一份 config。",
     italic=True, color=GREY, space_after=8)

heading("4.6 重排序 Reranker（新增）", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["LexicalReranker ★", "零依赖、中英文分词、可解释", "语义理解弱", "默认启用"],
        ["CrossEncoder(BGE) ★", "语义精排质量最高", "需 GPU/依赖重", "可选启用"],
        ["无重排", "最快", "top-k 精度低", "基线"],
        ["Cohere Rerank", "云端 SOTA", "需 API、付费", "云端可选"],
    ],
    widths=[1.4, 2.0, 2.0, 1.6],
)
para("结论：默认词法 Reranker 零成本提升精度；装 sentence-transformers 并设 RERANK_TYPE=cross_encoder 即可升级到 BGE 级精排（自动回退保底）。",
     italic=True, color=GREY, space_after=8)

heading("4.7 混合检索 Hybrid Search（新增）", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["向量 + BM25(RRF) ★", "兼顾语义与关键词、抗生僻词", "需维护词表", "已集成"],
        ["纯向量", "语义好", "关键词/专名易漏", "基线"],
        ["纯 BM25", "关键词准", "无语义", "基线"],
        ["ColBERT", "细粒度交互", "存储/算力大", "大规模再议"],
    ],
    widths=[1.5, 2.0, 1.9, 1.6],
)
para("结论：采用 RRF(k=60) 融合向量余弦与 BM25 得分，alpha 可调（0=纯向量，1=纯 BM25，默认 0.5）。",
     italic=True, color=GREY, space_after=8)

heading("4.8 对话记忆（新增）", level=2)
add_table(
    ["方案", "优势", "劣势", "本项目取舍"],
    [
        ["Token 预算 + 摘要 ★", "可控上下文、保留长期信息", "摘要消耗一次 LLM", "已采用"],
        ["最近 N 条(旧版)", "简单", "长对话丢信息", "已替换"],
        ["向量记忆检索", "相关性强", "复杂度高", "未来可选"],
        ["全量历史", "不丢信息", "易超窗口/费 token", "不采用"],
    ],
    widths=[1.5, 2.0, 1.9, 1.6],
)

heading("4.9 其他基础选型", level=2)
bullet("文本切分：自研段落感知切块器（按标题/空行分段，参数可调）；对比 LangChain RecursiveCharacterTextSplitter 更贴合中文文档，但缺句子边界感知。", bold_prefix="切分：")
bullet("配置：python-dotenv + 启动校验；对比 pydantic-settings 缺少类型校验，但足够轻量。", bold_prefix="配置：")
bullet("测试：pytest，36 用例覆盖去重、文件名解析、Agent 图构建、Reranker、Memory、混合检索等，均不触发真实 LLM/Chroma。", bold_prefix="测试：")

# ============================================================
# 第五章 关键技术机制
# ============================================================
heading("五、关键技术机制", level=1)

heading("5.1 Agent 多步循环（LangGraph ReAct）", level=2)
para("agent_answer() 入口构建 LangGraph 的 create_react_agent 状态图，max_iterations 映射为 recursion_limit。"
     "5 个工具用 @tool 装饰，自动生成 schema；通过 RunnableConfig 向工具注入 ToolContext，"
     "工具把检索 chunks / 待确认操作 / Observation 写入共享上下文——替代旧版全局可变状态，并发更安全。", space_after=6)
bullet("取得证据 → 用 rag_chain 专用合成器生成带引用答案；")
bullet("出现待确认操作 → 优先返回，绝不执行删除/重建；")
bullet("无检索（闲聊）→ 采用 Agent 自身自然语言回答；")
bullet("异常 → 安全回退；全程 build_trace 记录每一步 action/reason/input/output。")

heading("5.2 检索与生成解耦", level=2)
para("retrieve_tool 只返回原始 chunks，不直接调用 LLM；最终答案由 rag_chain.generate_answer_from_documents "
     "统一合成。好处：检索可被单独测试、可复用、生成策略可独立演进。", space_after=6)

heading("5.3 Query Rewrite", level=2)
para("检索前用 LLM 将口语化/指代模糊的问题改写为完整检索问句，提升召回；temperature=0 保证稳定。", space_after=6)

heading("5.4 Hybrid + Rerank 检索管线", level=2)
code_block(
    "用户问题\n"
    "  → rewrite_query()            # LLM 改写\n"
    "  → query_chunks()             # 向量召回（候选池 RERANK_TOP_N）\n"
    "  → get_all_chunks() + BM25     # 词法召回（hybrid 开启时）\n"
    "  → RRF 融合 (k=60, alpha 加权)\n"
    "  → reranker.rerank()          # 词法/CrossEncoder 精排\n"
    "  → top_k chunks → ToolContext"
)

heading("5.5 对话记忆管理", level=2)
para("memory.py 以 token 预算（默认 2000，中文按字、英文按词估算）做滑动窗口：预算内全保留；"
     "超预算时较早消息由 LLM 压缩为一段摘要（MEMORY_USE_SUMMARY=false 时退化为截断）。"
     "取代旧版“最近 8 条”硬编码，长对话不再丢上下文。", space_after=8)

# ============================================================
# 第六章 安全与配置
# ============================================================
heading("六、安全与配置", level=1)
para("上一版复盘指出的 P0 安全问题已全部修复并加固：", bold=True, space_after=4)
add_table(
    ["问题", "修复方式", "状态"],
    [
        [".env.example 明文 Key 泄露", "改为占位符 your-api-key-here，真实 Key 移至 .env", "已修复"],
        ["config.py 默认空格 Key 静默失败", "默认空串 + 启动 RuntimeError 校验", "已修复"],
        ["Key 进 git 历史", "用户轮换 Key；建议 git-filter-repo 清理（可选）", "已缓解"],
        [".env 变体可能误提交", ".gitignore 增加 .env.* 且 !.env.example 白名单", "已加固"],
        ["safe_filename 路径穿越", "移除 .. 与路径分隔符，仅保留文件名", "已修复"],
        ["危险操作无护栏", "pending_operation 二次确认 + 步数上限 + 回退", "原有能力保留"],
    ],
    widths=[2.2, 3.2, 1.0],
    font_size=9,
)
para("配置开关（config.py，可由 .env 覆盖）：HYBRID_SEARCH / HYBRID_ALPHA / RERANK_ENABLED / "
     "RERANK_TOP_N / RERANK_TYPE / MEMORY_MAX_TOKENS / MEMORY_USE_SUMMARY。",
     italic=True, color=GREY, space_after=8)

# ============================================================
# 第七章 测试与质量
# ============================================================
heading("七、测试与质量", level=1)
para("当前测试套件 36 个用例全部通过（pytest，~17s），覆盖不依赖真实 LLM/Chroma 的逻辑层：", space_after=4)
add_table(
    ["测试文件", "覆盖内容", "用例数"],
    [
        ["test_utils.py", "去重、工具函数", "4"],
        ["test_tools.py", "文件名解析（精确/歧义/未匹配）、工具上下文", "6"],
        ["test_agent.py", "LangGraph 图构建、工具注入、回退、trace", "6"],
        ["test_reranker.py", "词法/CrossEncoder 评分与重排", "5"],
        ["test_memory.py", "token 预算裁剪、摘要压缩、格式", "5"],
        ["test_document_loader.py", "多格式解析统一输出", "5"],
        ["test_retriever.py", "Hybrid 融合、Rerank 接入、动态配置", "5"],
    ],
    widths=[2.0, 3.2, 0.8],
    font_size=9,
)
para("额外保障：compileall 全量编译通过；app.py 可正常导入；build_agent_executor() 返回 CompiledStateGraph，"
     "5 个工具已正确接入。", space_after=8)

# ============================================================
# 第八章 升级历程
# ============================================================
heading("八、升级历程（本报告的来由）", level=1)
add_table(
    ["轮次", "内容", "主要产出"],
    [
        ["第 1 轮", "技术复盘", "HTML + Word 技术报告、7 项横向对比、问题清单"],
        ["第 2 轮", "安全修复(P0)", "清理 .env.example 明文 Key、config 启动校验、.gitignore 加固"],
        ["第 3 轮", "短期增强 + LangChain", "LLM 超时重试、流式输出、去重统一、Agent 升级为 LangGraph（18 用例）"],
        ["第 4 轮", "中期功能增强", "PyMuPDF、Reranker、Hybrid Search、记忆管理、更多格式（共 36 用例）"],
    ],
    widths=[0.8, 2.2, 3.0],
    font_size=9,
)

# ============================================================
# 第九章 代码质量评估
# ============================================================
heading("九、代码质量评估", level=1)
add_table(
    ["维度", "评分", "说明"],
    [
        ["架构清晰度", "9/10", "层次分明、依赖单向、职责单一"],
        ["模块化", "9/10", "16 模块，公共 util 去重，重复代码已消除"],
        ["可测试性", "8/10", "36 单测，逻辑层可独立测；LLM/Chroma 仍需集成测试"],
        ["安全性", "9/10", "P0 已修复、危险操作护栏、路径穿越修复"],
        ["可观测性", "8/10", "Agent trace 完整；缺结构化日志/指标"],
        ["生产就绪", "6/10", "单线程 UI、无并发、无鉴权、无 CI"],
        ["文档", "8/10", "README + 两份升级说明 + 本报告"],
    ],
    widths=[1.6, 1.0, 3.4],
    font_size=9,
)
para("综合：项目从原型已演进为工程化良好、可维护、有测试护航的 RAG Agent；"
     "主要短板在生产化（并发、鉴权、CI、可观测性）。", italic=True, color=GREY, space_after=8)

# ============================================================
# 第十章 后续建议
# ============================================================
heading("十、后续建议", level=1)
para("P0（建议尽快）：", bold=True, color=RGBColor(0xC0,0x39,0x2B), space_after=2)
bullet("将 Key 清理出 git 历史（git-filter-repo），即使已轮换也建议做；")
bullet("为 UI/API 增加基础鉴权，避免本地服务被局域网随意访问。")
para("P1（近期）：", bold=True, color=RGBColor(0xB9,0x77,0x0E), space_after=2)
bullet("引入 CI（GitHub Actions 跑 pytest + 编译检查）；")
bullet("为 Reranker 启用 BGE Cross-Encoder 实测精度收益；")
bullet("增加结构化日志与关键指标（检索耗时、命中率、回退率）。")
para("P2（可选增强）：", bold=True, color=ACCENT_LIGHT, space_after=2)
bullet("生产化：Streamlit → FastAPI + 前端，支持并发与流式 SSE；")
bullet("评估 Milvus/Qdrant（文档量进入百万级时）；")
bullet("表格/图片型 PDF 引入 pdfplumber / 多模态解析；")
bullet("向量记忆检索，进一步提升长程对话相关性。")

# ============================================================
# 附录
# ============================================================
heading("附录 A：依赖清单（requirements.txt 摘要）", level=1)
code_block(
    "streamlit  chromadb  pypdf  PyMuPDF  openai  python-dotenv\n"
    "langchain>=1.0  langchain-openai>=1.0  langchain-core>=1.0  langgraph>=0.3\n"
    "rank_bm25  python-docx  python-pptx  beautifulsoup4\n"
    "pytest\n"
    "# 可选：sentence-transformers（CrossEncoder 精排）· tiktoken（精确 token）"
)

heading("附录 B：新增/重构文件清单", level=1)
bullet("新增：src/reranker.py、src/memory.py、src/utils.py、tests/ 全套、两份升级说明")
bullet("重构：src/agent.py（→ LangGraph）、src/retriever.py（→ Hybrid+Rerank）、"
       "src/document_loader.py（→ 7 格式）、src/config.py（→ 开关+校验）、src/llm.py（→ 双轨+流式）")
bullet("加固：src/kb_operations.py（safe_filename）、.gitignore、.env.example")

doc.save(r"C:\Users\22234\Desktop\rag_kb_delete_update\技术复盘报告_最新版.docx")
print("OK: 技术复盘报告_最新版.docx generated")
